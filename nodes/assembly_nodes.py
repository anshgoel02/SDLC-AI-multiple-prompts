from __future__ import annotations

import json
from pathlib import Path

from ..debug import debug_state
from ..llm_utils import generate_text_with_usage, parse_json_model
from ..models import BRDAssemblyOut, BRDModel, BRDState
from ..template import build_template_spec

try:
    from docx import Document
except Exception:
    Document = None


def assembler_node(state: BRDState) -> BRDState:
    print("[node] assembler")
    template_spec = build_template_spec(Path(state.template_path))
    prompt = f"""SYSTEM
You are a BRD compiler and schema validator.

USER
Task:
1) Merge the ordered section drafts into a final BRDModel instance.
2) Also produce a single markdown document of the BRD (for stakeholders).
3) Normalize terms and deduplicate citations.
4) Add any remaining missing info to an “Open Questions” part.

Output:
Return valid JSON that matches this schema:
{{
  "brd_model": {{
    "sections": [string],
    "open_questions": [string]
  }},
  "brd_markdown": string
}}

INPUTS:
- Ordered sections (list of markdown strings): {json.dumps(state.section_drafts, indent=2)}
- FactPack: {json.dumps(state.facts.model_dump(), indent=2)}
- BRDTemplateSpec: {json.dumps(template_spec.model_dump(), indent=2)}
"""
    raw, usage = generate_text_with_usage(prompt, max_tokens=2000)
    if not raw.strip():
        print("[warn] assembler empty response; retrying once")
        retry_prompt = "Return ONLY valid JSON.\n\n" + prompt
        raw, usage = generate_text_with_usage(retry_prompt, max_tokens=2000)
    if not raw.strip():
        print("[warn] assembler still empty; using fallback markdown from section drafts")
        state.brd_model = BRDModel(
            sections=state.section_drafts,
            open_questions=[],
        )
        state.brd_markdown = "\n\n".join(state.section_drafts)
    else:
        data = parse_json_model(raw, BRDAssemblyOut)
        state.brd_model = data.brd_model
        state.brd_markdown = data.brd_markdown
    print(
        f"[node] assembler tokens_used={usage['total_tokens']} "
        f"(prompt={usage['prompt_tokens']} completion={usage['completion_tokens']})"
    )
    debug_state("assembler", state)
    return state


def persist_doc_node(state: BRDState) -> BRDState:
    print("[node] persist_doc")
    if not state.output_docx_path:
        print("[node] persist_doc tokens_used=0")
        debug_state("persist_doc", state)
        return state
    if Document is None:
        raise RuntimeError("python-docx is required to write .docx output.")
    doc = Document()
    for line in state.brd_markdown.splitlines():
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)
    doc.save(state.output_docx_path)
    print("[node] persist_doc tokens_used=0")
    debug_state("persist_doc", state)
    return state
